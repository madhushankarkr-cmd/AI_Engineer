import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")
    
if not my_api_key:
    raise ValueError("api error")

client = Groq(api_key = my_api_key)
model = "llama-3.3-70b-versatile"

job_description="""
Description
Role Overview
Do you want to build AI-powered developer tools and services that enable a billion builders to bring their ideas to life every day? Are you excited to build planet-scale platforms with enterprise trust, security, and reliability baked in from design to delivery? If so, this role in Microsoft’s CoreAI division is for you.

As a Software Engineer in our Developer Platforms team, you will design, build, and operate core platform services for developers across the entire application lifecycle. We build products and services for every need—from code management and CI/CD with Azure DevOps, to testing applications at scale with Azure App Testing, to running applications on Azure PaaS platforms such as Azure App Service, Container Apps, Functions, API Management, and Logic Apps.

We also build Azure- & Foundry-native integrations with leading technology providers (for example, MongoDB, Elastic, and Datadog) to give customers flexibility, choice, and confidence as they build and operate modern cloud applications—using AI as a core building block, not an afterthought.

Why CoreAI
Build trusted AI platforms and developer tools used by millions today—and billions tomorrow.
Solve deep technical problems at global, cloud-to-edge scale.
Apply AI to real-world engineering challenges, not just prototypes.
Grow in a culture that values engineering excellence, learning, and people.


Responsibilities
What You’ll Build
As part of a collaborative, platform-first engineering organization, you will build and evolve:

AI-powered cloud services and platforms that support developer and enterprise workflows, owning them end-to-end—from architecture and implementation to deployment and live-site operations.
Cloud-to-edge platform capabilities, including Azure resource providers, data-plane integrations, and portal experiences that enable secure, scalable management of modern applications and AI workloads.
AI-enabled engineering and lifecycle systems that improve testing efficiency, quality analysis, incident triage, and developer productivity across the software development lifecycle.
How You’ll Build and Grow
CoreAI combines deep engineering rigor, customer focus, and craftsmanship with a people-first learning culture, focused on building trusted AI platforms at scale.

Outcome-driven, platform-first approach: build durable platforms guided by metrics, telemetry, and customer feedback, with clear ownership across the full DevOps lifecycle.
AI-native engineering: design, build, and operate systems with AI embedded across design, coding, testing, release validation, and live-site operations, while maintaining correctness, safety, and trust.
Enterprise-grade trust: security, privacy, compliance, reliability, and responsible AI practices are treated as first-class engineering concerns.
Continuous growth and collaboration: deepen skills in distributed systems, cloud platforms, and applied AI through hands-on production work in an inclusive, supportive environment.


Qualifications
Required Qualifications
Bachelor’s or Master’s degree in Computer Science, or equivalent practical experience.
1-3 years of experience building production software using one or more modern programming languages such as C#, C++, Go, Java or Python.
Strong understanding of software engineering fundamentals, data structures, and problem-solving.
Ability to learn new technologies quickly and adapt to deliver customer and business impact.
Preferred Qualifications
Experience working in Linux environments and with open-source projects.
Familiarity with containers and orchestration technologies such as Docker and Kubernetes.
Experience with cloud infrastructure (Azure, AWS, or equivalent).
Exposure to site reliability engineering (SRE) practices.
Exposure to AI-assisted development and data-driven engineering workflows.
Knowledge of Azure resource providers, platform extensibility, and security, compliance, or responsible AI concepts.

This position will be open for a minimum of 5 days, with applications accepted on an ongoing basis until the position is filled.



Microsoft is an equal opportunity employer. All qualified applicants will receive consideration for employment without regard to age, ancestry, citizenship, color, family or medical care leave, gender identity or expression, genetic information, immigration status, marital status, medical condition, national origin, physical or mental disability, political affiliation, protected veteran or military status, race, ethnicity, religion, sex (including pregnancy), sexual orientation, or any other characteristic protected by applicable local laws, regulations and ordinances. If you need assistance with religious accommodations and/or a reasonable accommodation due to a disability during the application process, read more about requesting accommodations.
"""

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}

messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)


answer=response.choices[0].message.content

raw_json=answer
# print(raw_json)



import json
job_data=json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)



#parse real
class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    
    Return JSON matching this schema:
    {match_schema}

    Ensure the returned JSON structure is exactly:
    {{
        "score": <float representing overall match percentage from 0 to 100>,
        "details": {{
            "candidate_name": "<name of the candidate>",
            "matching_skills": [<list of matching skills>],
            "missing_important_skills": [<list of missing important skills>],
            "experience_requirement_met": <boolean indicating if experience requirement is met>,
            "final_verdict": "<short final verdict>"
        }}
    }}

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None



# lets do it now
resume_folder = Path(__file__).parent / "resumes"
all_results=[]
for file_path in resume_folder.iterdir():
    #C:\Users\Pratyush\padho_with_pratyush\week1\day5\resumes\abhay resume new - Abhay Singh.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    #score and details
    #acount chtgpt
    # request bhejna shhur krega millions
    #chattgot server jam ho jayega
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])